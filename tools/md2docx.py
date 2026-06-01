# -*- coding: utf-8 -*-
"""Convert the project's docs/*.md to docs/*.docx (one-shot helper)."""
import os, re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CHINESE_FONT = "Microsoft YaHei"
CODE_FONT = "Consolas"

def set_run_font(run, font_name=None, size=None, bold=None, italic=None, color=None):
    if font_name:
        run.font.name = font_name
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def add_inline_runs(p, text, font=CHINESE_FONT, base_bold=False):
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = p.add_run(text[pos:m.start()])
            set_run_font(run, font_name=font, bold=base_bold or None)
        token = m.group()
        if token.startswith("**") and token.endswith("**"):
            run = p.add_run(token[2:-2])
            set_run_font(run, font_name=font, bold=True)
        elif token.startswith("`") and token.endswith("`"):
            run = p.add_run(token[1:-1])
            set_run_font(run, font_name=CODE_FONT, color="C2632D",
                         bold=base_bold or None)
        elif token.startswith("[") and "](" in token:
            label = re.search(r'\[([^\]]+)\]', token).group(1)
            url = re.search(r'\(([^)]+)\)', token).group(1)
            run = p.add_run(label)
            set_run_font(run, font_name=font, color="0066CC",
                         bold=base_bold or None)
            run.font.underline = True
        pos = m.end()
    if pos < len(text):
        run = p.add_run(text[pos:])
        set_run_font(run, font_name=font, bold=base_bold or None)

def strip_inline_markers(text):
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text

def is_table_separator(line):
    return bool(re.match(r'^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$', line))

def parse_table_row(row):
    row = row.strip()
    if row.startswith("|"): row = row[1:]
    if row.endswith("|"): row = row[:-1]
    return [c.strip() for c in row.split("|")]

def md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Default font
    style = doc.styles['Normal']
    style.font.name = CHINESE_FONT
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), CHINESE_FONT)
    rFonts.set(qn('w:ascii'), CHINESE_FONT)
    rFonts.set(qn('w:hAnsi'), CHINESE_FONT)

    lines = content.split('\n')
    i, n = 0, len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Fenced code block
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # closing fence
            for cline in code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(cline if cline else " ")
                set_run_font(run, font_name=CODE_FONT, size=9, color="333333")
            continue

        # Heading
        m = re.match(r'^(#+)\s+(.+)$', stripped)
        if m:
            level = min(len(m.group(1)), 6)
            text = m.group(2)
            p = doc.add_heading(level=level)
            text_clean = strip_inline_markers(text)
            run = p.add_run(text_clean)
            set_run_font(run, font_name=CHINESE_FONT, bold=True)
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            quote_text = re.sub(r'^>\s*', '', stripped)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.7)
            add_inline_runs(p, quote_text, CHINESE_FONT)
            for r in p.runs:
                r.italic = True
                r.font.color.rgb = RGBColor.from_string("555555")
            i += 1
            continue

        # Table
        if "|" in stripped and i+1 < n and is_table_separator(lines[i+1]):
            header = parse_table_row(lines[i])
            i += 2
            rows = []
            while i < n and "|" in lines[i] and lines[i].strip():
                rows.append(parse_table_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1+len(rows), cols=len(header))
            try:
                table.style = 'Light Grid Accent 1'
            except KeyError:
                table.style = 'Table Grid'
            for j, cell_text in enumerate(header):
                cell = table.rows[0].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                add_inline_runs(p, cell_text, CHINESE_FONT, base_bold=True)
            for ri, row in enumerate(rows):
                for j, cell_text in enumerate(row):
                    if j >= len(header): break
                    cell = table.rows[ri+1].cells[j]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    add_inline_runs(p, cell_text, CHINESE_FONT)
            doc.add_paragraph()  # spacing after table
            continue

        # Unordered list
        if re.match(r'^[-*]\s+', stripped):
            text = re.sub(r'^[-*]\s+', '', stripped)
            try:
                p = doc.add_paragraph(style='List Bullet')
            except KeyError:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.7)
                text = "• " + text
            add_inline_runs(p, text, CHINESE_FONT)
            i += 1
            continue

        # Ordered list
        if re.match(r'^\d+\.\s+', stripped):
            text = re.sub(r'^\d+\.\s+', '', stripped)
            try:
                p = doc.add_paragraph(style='List Number')
            except KeyError:
                p = doc.add_paragraph()
            add_inline_runs(p, text, CHINESE_FONT)
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        add_inline_runs(p, stripped, CHINESE_FONT)
        i += 1

    doc.save(docx_path)
    print("Generated:", docx_path)

if __name__ == "__main__":
    import glob
    project = r"E:\yinji\yinji-project\yinji"
    # Scan all .md files under docs/ and devlog/ (including subfolders),
    # generate same-name .docx alongside each.
    for folder in ("docs", "devlog"):
        base = os.path.join(project, folder)
        if not os.path.isdir(base):
            continue
        for md in sorted(glob.glob(os.path.join(base, "**", "*.md"), recursive=True)):
            md_to_docx(md, md[:-3] + ".docx")
